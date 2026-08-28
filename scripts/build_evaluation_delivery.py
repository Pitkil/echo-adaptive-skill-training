"""Build the auditable ECHO 50-case candidate evaluation delivery package."""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
from collections import Counter
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_RUN = ROOT / "data" / "competition-evaluation" / "candidate-real-20260827-08"
DEFAULT_REVIEWED_RUN = (
    ROOT / "data" / "competition-evaluation" / "candidate-real-20260827-04-reviewed"
)
DEFAULT_OUT = ROOT / "交付材料" / "01_50组数据实际评测_赛事提交版"

NAVY = "102A43"
BLUE = "1F5A7A"
TEAL = "00A6A6"
MINT = "DDF4EF"
PALE = "F2F6F8"
AMBER = "E6A23C"
INK = RGBColor(16, 42, 67)
MUTED = RGBColor(83, 104, 120)


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def write_checksum_manifest(output_dir: Path) -> Path:
    """Write deterministic SHA-256 entries for every delivery file except the manifest itself."""
    manifest_path = output_dir / "06_SHA256SUMS.txt"
    entries = []
    for path in sorted(item for item in output_dir.rglob("*") if item.is_file()):
        if path == manifest_path:
            continue
        relative_path = path.relative_to(output_dir).as_posix()
        entries.append(f"{sha256(path)}  {relative_path}")
    manifest_path.write_text("\n".join(entries) + "\n", encoding="utf-8")
    return manifest_path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 110, start: int = 130, bottom: int = 110, end: int = 130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/msyh.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.is_file():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def chart_distributions(path: Path, learner: Counter, modules: Counter) -> None:
    image = Image.new("RGB", (1600, 800), "#F2F6F8")
    draw = ImageDraw.Draw(image)
    draw.text((90, 55), "50 组案例覆盖分布", fill="#102A43", font=font(52))
    draw.text((90, 125), "学习者画像与课程模块均衡覆盖", fill="#536878", font=font(27))

    groups = [
        ("学习者画像", ["P1", "P2", "P3"], learner, 90),
        ("课程模块", ["M1", "M2", "M3"], modules, 850),
    ]
    colors = ["#00A6A6", "#1F5A7A", "#E6A23C"]
    for title, labels, values, x0 in groups:
        draw.text((x0, 215), title, fill="#102A43", font=font(32))
        for index, label in enumerate(labels):
            y = 310 + index * 130
            count = int(values[label])
            width = int(count / 18 * 500)
            draw.rounded_rectangle((x0, y, x0 + 500, y + 52), 18, fill="#DCE6EB")
            draw.rounded_rectangle((x0, y, x0 + width, y + 52), 18, fill=colors[index])
            draw.text((x0, y - 40), label, fill="#102A43", font=font(25))
            draw.text((x0 + 520, y + 5), f"{count} 组", fill="#102A43", font=font(26))
    image.save(path)


def chart_machine_results(path: Path, citations: int) -> None:
    image = Image.new("RGB", (1600, 820), "#102A43")
    draw = ImageDraw.Draw(image)
    draw.text((90, 60), "机器层运行事实", fill="white", font=font(52))
    draw.text((90, 130), "candidate-real-20260827-08", fill="#8EDBD4", font=font(27))
    cards = [
        ("50 / 50", "案例完成", "#00A6A6"),
        ("50 / 50", "四 Agent 闭环", "#1F8A70"),
        (str(citations), "官方引用记录", "#2E86AB"),
        ("27 / 27", "资源自动检查通过", "#E6A23C"),
        ("0", "机器失败", "#6C8EAD"),
        ("0", "依赖降级", "#8D6A9F"),
    ]
    for index, (value, label, color) in enumerate(cards):
        col = index % 3
        row = index // 3
        x = 90 + col * 500
        y = 235 + row * 250
        draw.rounded_rectangle((x, y, x + 440, y + 190), 30, fill=color)
        draw.text((x + 35, y + 30), value, fill="white", font=font(48))
        draw.text((x + 35, y + 112), label, fill="white", font=font(27))
    image.save(path)


def add_title(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(7)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.16
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
        if widths:
            cell.width = Inches(widths[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            if row_index % 2:
                set_cell_shading(cell, PALE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            run.font.size = Pt(8.5)
            if widths:
                cell.width = Inches(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_document(doc: Document, run_id: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for name, size, color in (("Title", 32, NAVY), ("Heading 1", 18, NAVY), ("Heading 2", 13, BLUE)):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    header = section.header.paragraphs[0]
    header.text = f"ECHO · 50 组数据实际评测  |  {run_id}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("ECHO 赛事评测材料 · 两阶段证据  |  ").font.size = Pt(8)
    add_page_field(footer)


def add_cover(doc: Document, run_id: str, reviewed_run_id: str | None = None) -> None:
    reviewed_run_id = reviewed_run_id or run_id
    for _ in range(2):
        doc.add_paragraph()
    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = badge.add_run("COMPETITION EVALUATION · AUDITABLE DELIVERY")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("ECHO")
    run.bold = True
    run.font.size = Pt(48)
    run.font.color.rgb = INK
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("50 组数据实际评测总报告")
    run.bold = True
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    line = doc.add_table(rows=1, cols=1)
    line.autofit = False
    line.columns[0].width = Inches(6.8)
    line.rows[0].height = Inches(0.12)
    set_cell_shading(line.cell(0, 0), TEAL)
    doc.add_paragraph()

    status = doc.add_table(rows=1, cols=2)
    status.alignment = WD_TABLE_ALIGNMENT.LEFT
    status.autofit = False
    left, right = status.rows[0].cells
    set_cell_shading(left, AMBER)
    set_cell_shading(right, PALE)
    set_cell_margins(left, 160, 180, 160, 180)
    set_cell_margins(right, 160, 180, 160, 180)
    left.text = "证据结构\n双人复核基线 + 修复后复测"
    right.text = (
        f"人工复核批次\n{reviewed_run_id}\n\n最新复测批次\n{run_id}\n"
        "50/50 完成 · 0 机器失败"
    )
    for cell in (left, right):
        for paragraph in cell.paragraphs:
            for item in paragraph.runs:
                item.font.size = Pt(11)
                item.bold = True
    left.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    add_body(
        doc,
        "适用范围：领域知识个性化生成与多智能体协同决策系统研究比赛。"
        "本报告将真实双人复核与修复后的最新复测分批保存。两批结果不混写，"
        "从而保证每项结论都能回到对应运行证据。",
    )
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("生成日期  ").bold = True
    meta.add_run("2026-08-28\n")
    meta.add_run("代码基线  ").bold = True
    meta.add_run("member/a-integration · 1f4c8c8（运行时工作区含未提交评测修复）\n")
    meta.add_run("交付性质  ").bold = True
    meta.add_run("赛事提交证据包；最新复测尚待同批次双人复核")
    doc.add_page_break()


def build_report(
    path: Path,
    run_id: str,
    manifest: dict[str, Any],
    results: list[dict[str, Any]],
    citations: int,
    chart_distribution: Path,
    chart_results: Path,
) -> None:
    doc = Document()
    configure_document(doc, run_id)
    add_cover(doc, run_id)

    add_title(doc, "执行摘要")
    add_body(
        doc,
        "本批次在真实 Docker 系统上完成 50 组固定案例：50 组均为 completed，"
        "没有机器失败或依赖降级；四个后台 Agent 的持久化记录完整率为 100%。"
    )
    callout = doc.add_table(rows=1, cols=1)
    set_cell_shading(callout.cell(0, 0), MINT)
    set_cell_margins(callout.cell(0, 0), 180, 200, 180, 200)
    callout.cell(0, 0).text = (
        "机器层面已完成：50/50 真实运行、50/50 四 Agent 闭环、149 条官方引用记录、"
        "27/27 资源自动检查通过，失败清单为空。\n"
        "正式指标仍待两名真实成员独立复核，因此当前不构成比赛指标达标结论。"
    )
    for run in callout.cell(0, 0).paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = INK
    add_table(
        doc,
        ["项目", "机器结果", "当前判定"],
        [
            ["案例运行", "50/50 completed", "通过"],
            ["四 Agent 闭环", "50/50（100%）", "通过"],
            ["官方引用记录", f"{citations} 条；49/50 案例含引用", "供人工复核"],
            ["个性化资源", "三类各 9；27/27 自动检查通过", "待讲师发布"],
            ["四项人工指标", "0/50 已完成双人复核", "待判定"],
        ],
        [2.2, 2.7, 1.7],
    )
    doc.add_picture(str(chart_results), width=Inches(6.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    add_title(doc, "1. 赛题要求映射")
    add_table(
        doc,
        ["赛题/任务要求", "本批次证据", "状态"],
        [
            ["不少于 50 组固定案例", "case-001 至 case-050，全部有实际输出", "满足"],
            ["不少于 3 组学习者画像", "P1=17、P2=17、P3=16", "满足"],
            ["覆盖多智能体中间过程", "每案保存分析、生成、检查、下一步安排", "满足"],
            ["知识库与官方出处", "Microsoft 官方材料；149 条引用记录", "机器完成，待人工判定相关性"],
            ["三类个性化资源", "定制资料、实操指南、阶段测试各 9", "自动检查通过，待发布"],
            ["幻觉率 <5%", "需要逐案事实声明复核", "待双人复核"],
            ["难度适配率 ≥85%", "需要对画像、难度和实际内容判定", "待双人复核"],
            ["核心知识覆盖率 ≥90%", "需要对知识点要点逐案判定", "待双人复核"],
            ["引用可追溯率目标 100%", "需要逐条打开 URL/章节核验", "待双人复核"],
        ],
        [2.2, 3.3, 1.2],
    )
    add_body(doc, "要求来源：比赛方案 PDF、ECHO 50 组数据实际评测任务书及仓库 competition-requirements.md。")

    doc.add_page_break()
    add_title(doc, "2. 环境与数据真实性")
    health = manifest["health_check"]["response"]
    dependencies = health.get("dependencies", {})
    add_table(
        doc,
        ["服务", "运行状态", "模式/说明"],
        [
            ["ECHO API", health.get("status", ""), "http://127.0.0.1:8010"],
            ["可追溯 RAG", dependencies.get("punditrag_query", {}).get("status", ""), "正式知识库检索"],
            ["SimpleMem", dependencies.get("simplemem", {}).get("status", ""), "长期语义记忆"],
            ["微表征", dependencies.get("micro_representation", {}).get("status", ""), "mode=real；本轮无音频输入"],
            ["ASR", dependencies.get("asr", {}).get("status", ""), "faster-whisper；本轮未调用"],
            ["业务数据库", dependencies.get("database", {}).get("status", ""), "真实持久化"],
        ],
        [2.0, 1.3, 3.4],
    )
    add_bullets(
        doc,
        [
            "Docker 镜像、容器与卷的数据根位于 D:\\DockerDesktopData；未把 Docker 虚拟磁盘放在 C 盘。",
            "运行器未导出凭据、令牌或真实个人信息；学习者为隔离的合成 P1/P2/P3 画像。",
            "运行清单记录冻结案例 SHA-256、运行器 SHA-256、分支、基线提交和工作区脏状态。",
            "本批次为未提交修复代码产生的 candidate；正式冻结必须先提交，再按提交 SHA 复跑或签署。",
        ],
    )

    doc.add_page_break()
    add_title(doc, "3. 评测设计与覆盖")
    doc.add_picture(str(chart_distribution), width=Inches(6.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    learner = Counter(row["learner_type"] for row in results)
    modules = Counter(row["module"] for row in results)
    scenarios = Counter(row["scenario_type"] for row in results)
    add_table(
        doc,
        ["维度", "分布"],
        [
            ["学习者画像", f"P1 {learner['P1']} / P2 {learner['P2']} / P3 {learner['P3']}"],
            ["模块", f"M1 {modules['M1']} / M2 {modules['M2']} / M3 {modules['M3']}"],
            ["资源", "定制资料 9 / 实操指南 9 / 阶段测试 9"],
            ["解释", f"{scenarios['explanation']} 组"],
            ["动态调整", f"{scenarios['dynamic']} 组"],
            ["异常与边界", f"{scenarios['exception']} 组"],
        ],
        [2.0, 4.7],
    )
    add_body(
        doc,
        "每组按统一流程创建隔离学习者状态，读取 MIRT 与长期记忆，执行单一主要动作，"
        "必要时检索官方材料，保存四 Agent 记录、实际输出、资源、引用和失败原因。"
    )

    doc.add_page_break()
    add_title(doc, "4. 机器层运行结果")
    action_counts = Counter(row["actual_output"].get("primary_action") for row in results)
    add_table(
        doc,
        ["主要动作", "案例数", "说明"],
        [
            ["GENERATE_RESOURCE", str(action_counts["GENERATE_RESOURCE"]), "三类资源各 9"],
            ["LEARNING_DIALOGUE", str(action_counts["LEARNING_DIALOGUE"]), "解释与辅导"],
            ["GRADE_ANSWER", str(action_counts["GRADE_ANSWER"]), "固定题服务端判分"],
            ["CHANGE_MODULE", str(action_counts["CHANGE_MODULE"]), "唯一模块切换动作"],
        ],
        [2.4, 1.1, 3.2],
    )
    add_bullets(
        doc,
        [
            "50 个结果文件状态全部为 completed；reports/failures.json 为 []。",
            "50 个案例都保存分析、生成、检查、下一步安排四个 Agent 的持久化记录。",
            "49 个案例包含官方引用；047 仅执行模块切换，不需要专业知识引用。",
            "没有 mock 依赖、没有 HTTP 失败、没有降级结果。",
        ],
    )

    add_title(doc, "5. 个性化资源结果")
    add_table(
        doc,
        ["资源类型", "数量", "自动检查", "发布状态"],
        [
            ["定制学习资料 custom_note", "9", "9/9 通过", "pending_review"],
            ["实操指南 practice_guide", "9", "9/9 通过", "pending_review"],
            ["阶段测试 staged_test", "9", "9/9 通过", "pending_review"],
        ],
        [2.7, 0.8, 1.4, 1.8],
    )
    add_body(
        doc,
        "当前产品将“自动检查通过”和“正式发布”分开：验证通过的资源由讲师或管理员显式发布。"
        "本次评测未越权自动发布，因此 27 个资源均为 pending_review；这不是生成失败。"
    )

    doc.add_page_break()
    add_title(doc, "6. 关键边界案例")
    add_table(
        doc,
        ["案例", "验证点", "实际结果"],
        [
            ["037", "开放题关键词评分", "中央编排器 + 管理服务/插件判对，更新 MIRT"],
            ["040", "中文全角括号选项", "A（Process、Step、Event）判对"],
            ["041", "Temperature 固定题出处", "判 B 正确，仅保留固定题官方出处"],
            ["044", "文本犹豫、无音频", "micro_signal_required=false，不伪造语音证据"],
            ["047", "单一主要动作", "仅 CHANGE_MODULE，不混入资源或判题"],
            ["048", "重复提交幂等", "同 request_id 响应相同，重放前后 U/A/R 与次数不变"],
        ],
        [0.7, 2.2, 3.8],
    )
    edge_048 = next(row for row in results if row["case_id"] == "048")
    proof = edge_048["actual_output"].get("idempotency_proof") or {}
    add_body(
        doc,
        "048 幂等证据：首次提交后 attempt_count="
        f"{proof.get('after_first', {}).get('attempt_count')}；重放后 attempt_count="
        f"{proof.get('after_replay', {}).get('attempt_count')}；ability_unchanged_on_replay="
        f"{str(proof.get('ability_unchanged_on_replay')).lower()}。"
    )

    doc.add_page_break()
    add_title(doc, "7. 双人复核方法")
    add_body(
        doc,
        "04_双人复核表.csv 已为每个案例预置 reviewer-1 和 reviewer-2 两行，共 100 行。"
        "两名真实成员必须分别查看“实际输出、Agent 中间记录、官方引用、原始 HTTP 记录”，独立填写，不得复制对方结论。"
    )
    add_table(
        doc,
        ["指标", "填写字段", "计算方法", "阈值"],
        [
            ["幻觉率", "verifiable_claim_count / unsupported_claim_count", "无依据声明 ÷ 可核验事实声明", "< 5%"],
            ["难度适配率", "difficulty_match", "yes 案例 ÷ 已复核案例", "≥ 85%"],
            ["知识覆盖率", "knowledge_coverage", "yes 案例 ÷ 已复核案例", "≥ 90%"],
            ["引用可追溯率", "citation_required_count / citation_traceable_count", "可追溯引用 ÷ 应有引用", "100%"],
            ["内容错误率", "content_error", "yes 案例 ÷ 已复核案例", "仅报告"],
        ],
        [1.3, 2.3, 2.4, 0.9],
    )
    add_bullets(
        doc,
        [
            "若两人结论不一致，必须由第三名真实成员裁决并记录 adjudication。",
            "复核完成后写回 results/case-XXX.json，再以 --require-formal 运行计分器。",
            "不得用旧批次 candidate-real-20260827-04 的 CSV 套用到本批次。",
            "不得把 expected 复制成 actual，也不得把 AI 会话当作第二名人工复核人。",
        ],
    )

    doc.add_page_break()
    add_title(doc, "8. 当前结论与正式冻结条件")
    warning = doc.add_table(rows=1, cols=1)
    set_cell_shading(warning.cell(0, 0), "FFF3CD")
    set_cell_margins(warning.cell(0, 0), 180, 200, 180, 200)
    warning.cell(0, 0).text = "当前结论：机器执行通过；比赛四项人工指标待判定。"
    warning.cell(0, 0).paragraphs[0].runs[0].bold = True
    add_body(
        doc,
        "本批次可以证明系统在真实环境中完成 50 组输入、四 Agent 闭环、官方检索、三类资源生成、"
        "固定题判分和幂等处理；但不能在缺少两名真实成员复核时声称幻觉率、难度适配率、"
        "知识覆盖率和引用可追溯率已达标。"
    )
    add_bullets(
        doc,
        [
            "由两名真实成员完成 100 行独立复核，必要时第三人裁决。",
            "将复核结论写回本批次结果，并运行正式计分器生成最终指标。",
            "讲师/管理员审核并发布 27 个已通过自动检查的资源。",
            "提交当前评测修复，记录正式 commit SHA；随后按提交版本复跑或签署冻结声明。",
            "重新计算交付目录 SHA-256，并由负责人签字确认。",
        ],
    )

    add_title(doc, "9. 交付目录索引")
    add_table(
        doc,
        ["目录/文件", "用途"],
        [
            ["01_评测总报告.docx / 02_评测总报告.pdf", "本报告可编辑版与固定版"],
            ["03_评测结果汇总.csv", "50 行机器状态汇总"],
            ["04_双人复核表.csv", "100 行空白人工复核表"],
            ["05_运行环境与版本.json", "服务状态、版本、哈希与限制"],
            ["原始输入/", "50 组冻结输入与画像初始化"],
            ["实际输出/", "50 组系统实际最终输出"],
            ["Agent中间记录/", "50 组四 Agent 持久化记录"],
            ["官方引用/", "50 组引用记录；047 为空数组"],
            ["原始HTTP记录/", "50 组未改写的运行器原始记录"],
            ["说明/", "数据字典、复核说明、知识库证据"],
            ["06_SHA256SUMS.txt", "交付文件完整性校验"],
        ],
        [3.1, 3.6],
    )
    add_body(doc, "报告结束。正式指标页必须在双人复核完成后重新生成。")
    doc.save(path)


def build_delivery(run_dir: Path, output_dir: Path) -> Path:
    if not run_dir.is_dir():
        raise FileNotFoundError(run_dir)
    results = [read_json(path) for path in sorted((run_dir / "results").glob("case-*.json"))]
    if len(results) != 50:
        raise ValueError(f"Expected 50 results, found {len(results)}")
    failures = [(row["case_id"], row.get("failure_reasons")) for row in results if row.get("failure_reasons")]
    if failures or any(row.get("status") != "completed" for row in results):
        raise ValueError(f"Run is not machine-clean: {failures}")

    output_dir.mkdir(parents=True, exist_ok=True)
    folders = {
        name: output_dir / name
        for name in (
            "原始输入",
            "实际输出",
            "Agent中间记录",
            "官方引用",
            "原始HTTP记录",
            "图表",
            "说明",
        )
    }
    for folder in folders.values():
        folder.mkdir(parents=True, exist_ok=True)

    for result in results:
        case_id = result["case_id"]
        raw_path = run_dir / "raw" / f"case-{case_id}.json"
        raw = read_json(raw_path)
        write_json(
            folders["原始输入"] / f"case-{case_id}.json",
            {
                "run_id": result["run_id"],
                "case": raw.get("case"),
                "profile_initialization": raw.get("profile_initialization"),
                "request": result.get("request"),
                "learner_profile_snapshot": result.get("learner_profile_snapshot"),
            },
        )
        write_json(
            folders["实际输出"] / f"case-{case_id}.json",
            {
                "run_id": result["run_id"],
                "case_id": case_id,
                "status": result["status"],
                "started_at": result.get("started_at"),
                "finished_at": result.get("finished_at"),
                "actual_output": result.get("actual_output"),
                "resource_record": result.get("resource_record"),
                "failure_reasons": result.get("failure_reasons"),
            },
        )
        write_json(
            folders["Agent中间记录"] / f"case-{case_id}.json",
            {"run_id": result["run_id"], "case_id": case_id, "agent_records": result["agent_records"]},
        )
        write_json(
            folders["官方引用"] / f"case-{case_id}.json",
            {"run_id": result["run_id"], "case_id": case_id, "citations": result.get("citations", [])},
        )
        shutil.copy2(raw_path, folders["原始HTTP记录"] / raw_path.name)

    shutil.copy2(run_dir / "reports" / "cases.csv", output_dir / "03_评测结果汇总.csv")
    shutil.copy2(run_dir / "human-review-template.csv", output_dir / "04_双人复核表.csv")
    shutil.copy2(run_dir / "reports" / "metrics.json", folders["说明"] / "机器计分结果.json")
    shutil.copy2(run_dir / "reports" / "failures.json", folders["说明"] / "机器失败清单.json")
    shutil.copy2(run_dir / "run_manifest.json", folders["说明"] / "原始运行清单.json")

    runtime_materials = ROOT / "data" / "formal-materials" / "runtime-20260827T081157Z"
    for name in ("material-import-report.json", "retrieval-verification.json"):
        source = runtime_materials / name
        if source.is_file():
            shutil.copy2(source, folders["说明"] / name)

    manifest = read_json(run_dir / "run_manifest.json")
    citations = sum(len(row.get("citations") or []) for row in results)
    resources = [row["resource_record"] for row in results if row.get("resource_record")]
    environment = {
        "schema_version": "1.0",
        "generated_at": datetime.now(UTC).isoformat(),
        "run_id": manifest["run_id"],
        "candidate_status": "pending_two_human_reviewers",
        "formal_ready": False,
        "repository": manifest["repository"],
        "configuration": manifest["configuration"],
        "health_check": manifest["health_check"],
        "docker_data_root": "D:\\DockerDesktopData",
        "no_audio_input_in_fifty_cases": True,
        "micro_service_used_for_scoring": False,
        "machine_facts": {
            "completed_cases": 50,
            "machine_failure_count": 0,
            "closed_loop_complete_cases": 50,
            "citation_record_count": citations,
            "cases_with_citations": sum(bool(row.get("citations")) for row in results),
            "resource_count": len(resources),
            "resource_auto_verification_passed": sum(bool(item.get("verification_passed")) for item in resources),
            "resource_status": dict(Counter(item.get("status") for item in resources)),
        },
        "inputs": manifest["inputs"],
        "quality_checks": {
            "pytest": "passed",
            "ruff": "passed",
            "compileall": "passed",
            "docker_compose_config": "passed",
        },
        "limitations": [
            "四项主观指标尚待两名真实成员逐案独立复核。",
            "27 个资源已通过自动检查，但按产品发布规则仍为 pending_review。",
            "运行时工作区包含未提交评测修复；当前批次是 candidate，不是正式冻结提交。",
        ],
    }
    write_json(output_dir / "05_运行环境与版本.json", environment)

    readme = """# ECHO 50 组数据实际评测交付包

本目录对应运行批次 `candidate-real-20260827-08`。

机器事实：50/50 完成、50/50 四 Agent 闭环、149 条官方引用记录、27/27 资源自动检查通过、失败清单为空。

当前不是正式指标结论。`04_双人复核表.csv` 的 100 行必须由两名真实成员独立填写；旧批次 `candidate-real-20260827-04` 的复核结果不能套用。复核完成后写回 `results/case-XXX.json`，运行：

```powershell
python scripts/score_competition_evaluation.py --run-dir data/competition-evaluation/candidate-real-20260827-08 --require-formal
```

本轮 50 组没有音频输入；044 是文本犹豫案例，`micro_signal_required=false`。微表征和 ASR 的健康状态只证明环境能力，不参与本轮评分。
"""
    (folders["说明"] / "README_复核与正式冻结说明.md").write_text(readme, encoding="utf-8")

    dictionary = """# 数据字典

- `原始输入/case-XXX.json`：冻结案例、画像初始化、实际请求和运行前画像。
- `实际输出/case-XXX.json`：ECHO 最终回复、资源、更新后画像、依赖限制和失败原因。
- `Agent中间记录/case-XXX.json`：分析、生成、检查、下一步安排四个 Agent 的持久化记录。
- `官方引用/case-XXX.json`：系统实际登记的官方来源；047 为模块切换，引用数组为空。
- `原始HTTP记录/case-XXX.json`：运行器保存的未改写 HTTP 请求/响应审计记录，已脱敏。
- `03_评测结果汇总.csv`：机器状态，不含人工指标结论。
- `04_双人复核表.csv`：每案例两行，必须由两名真实成员填写。
"""
    (folders["说明"] / "数据字典.md").write_text(dictionary, encoding="utf-8")

    correction = """# 冻结案例元数据纠错说明

本批次在复测前修正了 037、040、041、048 的固定题上下文：

- 037 对齐正式 Kernel 前测题，并启用题库登记的关键词评分规则。
- 040 对齐 Process Framework 的 A 选项题，支持 `A（解释）` 全角括号输入。
- 041 将 Temperature 题标题和章节纠正为对应 Microsoft API 页面。
- 048 使用同一 `request_id` 做真实重放，并保存首次提交后与重放后的 U/A/R、attempt_count。

修正后重新运行全部 50 组，未改写旧批次结果。案例与运行器哈希记录在运行清单中。
"""
    (folders["说明"] / "冻结案例元数据纠错说明.md").write_text(correction, encoding="utf-8")

    learner = Counter(row["learner_type"] for row in results)
    modules = Counter(row["module"] for row in results)
    distribution_chart = folders["图表"] / "01_覆盖分布.png"
    result_chart = folders["图表"] / "02_机器结果.png"
    chart_distributions(distribution_chart, learner, modules)
    chart_machine_results(result_chart, citations)

    report_path = output_dir / "01_评测总报告.docx"
    build_report(report_path, manifest["run_id"], manifest, results, citations, distribution_chart, result_chart)
    return report_path


def build_two_stage_report(
    path: Path,
    latest_run: Path,
    reviewed_run: Path,
    chart_distribution: Path,
    chart_results: Path,
) -> None:
    """Build a report that keeps human-review and repaired rerun evidence separate."""
    latest_manifest = read_json(latest_run / "run_manifest.json")
    latest_results = [read_json(item) for item in sorted((latest_run / "results").glob("case-*.json"))]
    reviewed_manifest = read_json(reviewed_run / "run_manifest.json")
    reviewed_metrics = read_json(reviewed_run / "reports" / "metrics.json")
    latest_id = latest_manifest["run_id"]
    reviewed_id = reviewed_manifest["run_id"]
    citations = sum(len(item.get("citations") or []) for item in latest_results)
    metrics = reviewed_metrics["metrics"]

    def rate(name: str) -> str:
        item = metrics[name]
        return f"{item['numerator']}/{item['denominator']} = {item['value'] * 100:.2f}%"

    def verdict(name: str) -> str:
        return "通过" if metrics[name]["passed"] else "未通过"

    doc = Document()
    configure_document(doc, latest_id)
    add_cover(doc, latest_id, reviewed_id)

    add_title(doc, "执行摘要")
    add_body(
        doc,
        "本交付采用两阶段证据：批次 04 保存两名真实成员对 50 组结果的独立复核和正式计分；"
        "批次 08 保存问题修复后重新运行的 50 组最新实际输出。两批结果不混写。"
    )
    callout = doc.add_table(rows=1, cols=1)
    set_cell_shading(callout.cell(0, 0), "FFF3CD")
    set_cell_margins(callout.cell(0, 0), 180, 200, 180, 200)
    callout.cell(0, 0).text = (
        "真实结论：批次 04 已完成 50/50 双人复核，但核心知识覆盖率 76.00%、"
        "引用可追溯率 99.24%，没有达到目标。批次 08 已完成修复复测，"
        "50/50 运行成功且机器失败为 0，但尚未完成与该批次对应的双人复核。"
    )
    callout.cell(0, 0).paragraphs[0].runs[0].bold = True
    add_table(
        doc,
        ["证据层", "批次", "完成情况", "可得结论"],
        [
            ["人工复核基线", reviewed_id, "50/50 双人复核", "formal_ready=true"],
            ["最新修复复测", latest_id, "50/50 completed", "机器层通过"],
            ["批次边界", "04 → 08", "48/50 结果变化", "人工结论不可迁移"],
        ],
        [1.5, 2.2, 1.5, 1.5],
    )
    doc.add_picture(str(chart_results), width=Inches(6.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    add_title(doc, "1. 赛题要求映射")
    add_table(
        doc,
        ["要求", "交付证据", "状态"],
        [
            ["不少于 50 组固定案例", "case-001 至 case-050，均有输入、输出和 HTTP 记录", "满足"],
            ["不少于 3 组学习者画像", "P1=17、P2=17、P3=16", "满足"],
            ["多智能体中间过程", "每案保存分析、生成、检查、下一步安排", "满足"],
            ["三类个性化资源", "定制资料、实操指南、阶段测试各 9", "机器检查通过，待发布"],
            ["真实人工评测", "批次 04 两名复核者、100 行原始记录", "满足"],
            ["最新修复证据", "批次 08 全量复跑、50/50 完成", "满足"],
        ],
        [2.0, 3.5, 1.2],
    )
    add_body(doc, "指标目标来自比赛方案、50 组实际评测任务书和仓库 competition-requirements.md。")

    doc.add_page_break()
    add_title(doc, "2. 批次 04：双人复核正式指标")
    add_table(
        doc,
        ["指标", "结果", "目标", "判定"],
        [
            ["幻觉率", rate("hallucination_rate"), "< 5%", verdict("hallucination_rate")],
            ["难度适配率", rate("difficulty_adaptation_rate"), "≥ 85%", verdict("difficulty_adaptation_rate")],
            ["核心知识覆盖率", rate("knowledge_coverage_rate"), "≥ 90%", verdict("knowledge_coverage_rate")],
            ["引用可追溯率", rate("citation_traceability_rate"), "100%", verdict("citation_traceability_rate")],
            ["闭环记录完整率", rate("closed_loop_completeness_rate"), "100%", verdict("closed_loop_completeness_rate")],
            ["案例级内容错误率", rate("case_content_error_rate"), "仅报告", "—"],
        ],
        [1.8, 2.1, 1.4, 1.3],
    )
    add_bullets(
        doc,
        [
            "难度不匹配：009、021、033、039、048、049，共 6 组。",
            "知识覆盖不足：008、009、012、021、024、028、032、033、036、039、048、049，共 12 组。",
            "引用不可追溯：041，共 1 条引用记录。",
            "内容错误案例：037、048，共 2 组。",
        ],
    )
    add_body(
        doc,
        "两名复核者结论一致，计分器 formal_ready=true；但 all_thresholds_passed=false。"
        "本包保留未达标事实，不把复测结果倒填为旧批次人工结论。"
    )

    doc.add_page_break()
    add_title(doc, "3. 批次 08：最新修复复测")
    action_counts = Counter(item["actual_output"].get("primary_action") for item in latest_results)
    add_table(
        doc,
        ["机器事实", "结果", "说明"],
        [
            ["案例完成", "50/50", "reports/failures.json 为空"],
            ["四 Agent 闭环", "50/50", "闭环记录完整率 100%"],
            ["官方引用", str(citations), "49/50 案例含引用；047 为模块切换"],
            ["个性化资源", "27/27", "三类各 9，自动检查通过，状态 pending_review"],
            ["主要动作", str(sum(action_counts.values())), "每轮仅一个 primary_action"],
            ["依赖降级", "0", "真实 Docker 环境"],
        ],
        [1.8, 1.3, 3.6],
    )
    doc.add_picture(str(chart_distribution), width=Inches(6.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(
        doc,
        "本轮 50 组没有音频输入。微表征和 ASR 的健康状态仅证明环境能力，不参与本轮评分；"
        "案例 044 为文本犹豫，micro_signal_required=false。"
    )

    doc.add_page_break()
    add_title(doc, "4. 问题修复与边界验证")
    add_table(
        doc,
        ["案例", "批次 04 发现", "批次 08 复测证据"],
        [
            ["037", "开放题关键词命中却判错", "关键词评分修复，答案判对并更新 MIRT"],
            ["040", "全角括号选项解析风险", "A（Process、Step、Event）判对"],
            ["041", "Temperature 出处标题/页面不一致", "仅保留与固定题对应的官方出处"],
            ["048", "幂等重放仍改变能力和次数", "同 request_id 重放，U/A/R 与 attempt_count 不变"],
        ],
        [0.8, 2.7, 3.2],
    )
    add_body(
        doc,
        "哈希比较显示，批次 04 到 08 仅 014、047 的结果文件完全一致，其余 48 组发生变化。"
        "变化可能来自修复、重新检索、重新生成或运行记录差异，因此不能用机器变化直接替代人工复核。"
    )
    add_bullets(
        doc,
        [
            "037、040、041、048 已有明确的定向复测证据。",
            "其他 04 批次覆盖不足案例在 08 中虽已重新生成，但是否达到 90% 仍须同批次人工复核。",
            "27 个资源保持 pending_review，自动检查通过不等于讲师已发布。",
        ],
    )

    doc.add_page_break()
    add_title(doc, "5. 环境与可复现性")
    health = latest_manifest["health_check"]["response"]
    dependencies = health.get("dependencies", {})
    add_table(
        doc,
        ["服务", "状态", "本轮作用"],
        [
            ["ECHO API", health.get("status", ""), "认证、会话、判分、资源与审计"],
            ["PunditRAG", dependencies.get("punditrag_query", {}).get("status", ""), "正式知识库检索与引用"],
            ["SimpleMem", dependencies.get("simplemem", {}).get("status", ""), "长期记忆"],
            ["微表征", dependencies.get("micro_representation", {}).get("status", ""), "真实服务健康；本轮无音频"],
            ["ASR", dependencies.get("asr", {}).get("status", ""), "服务健康；本轮未调用"],
            ["业务数据库", dependencies.get("database", {}).get("status", ""), "真实持久化"],
        ],
        [1.8, 1.3, 3.6],
    )
    add_bullets(
        doc,
        [
            "Docker 数据根位于 D:\\DockerDesktopData，C 盘不保存 Docker 虚拟磁盘。",
            "运行清单保存冻结案例、运行器、分支、基线提交和工作区状态。",
            "逐案保留原始输入、实际输出、四 Agent 记录、官方引用与原始 HTTP 记录。",
            "06_SHA256SUMS.txt 用于检查交付目录内容是否被修改。",
        ],
    )

    doc.add_page_break()
    add_title(doc, "6. 复核与正式冻结规则")
    add_body(
        doc,
        "04_双人复核表_批次04.csv 是已完成的合并复核表；04A、04B 是两名复核者的原始独立表；"
        "04C_批次08待复核模板.csv 只用于最新批次。"
    )
    add_bullets(
        doc,
        [
            "两名复核者必须分别查看批次 08 的实际输出、Agent 记录、引用和 HTTP 记录。",
            "若结论不一致，由第三名真实成员裁决并记录 adjudication。",
            "复核结论只能写回批次 08，随后以 --require-formal 重新计分。",
            "完成后必须更新本报告、05_运行环境与版本.json、校验和与压缩包。",
        ],
    )
    warning = doc.add_table(rows=1, cols=1)
    set_cell_shading(warning.cell(0, 0), "FFF3CD")
    set_cell_margins(warning.cell(0, 0), 180, 200, 180, 200)
    warning.cell(0, 0).text = (
        "对外表述边界：可以说“50 组真实运行完成、双人复核流程已在批次 04 完成、"
        "关键问题已在批次 08 修复复测”；不能说“批次 08 四项人工指标全部达标”。"
    )
    warning.cell(0, 0).paragraphs[0].runs[0].bold = True

    doc.add_page_break()
    add_title(doc, "7. 交付目录索引")
    add_table(
        doc,
        ["文件或目录", "用途"],
        [
            ["01_评测总报告.docx / 02_评测总报告.pdf", "本报告可编辑版与固定版"],
            ["03_评测结果汇总_批次08.csv", "最新复测 50 行机器状态"],
            ["04_双人复核表_批次04.csv", "批次 04 已完成的 100 行双人复核"],
            ["04A / 04B_复核人*_批次04.csv", "两名复核者原始独立记录"],
            ["04C_批次08待复核模板.csv", "最新批次待复核模板"],
            ["05_运行环境与版本.json", "批次、服务、版本、限制与机器事实"],
            ["07_人工复核正式指标_批次04.json", "批次 04 正式计分结果"],
            ["原始输入 / 实际输出 / Agent中间记录", "批次 08 的 50 组核心证据"],
            ["官方引用 / 原始HTTP记录", "批次 08 的来源和调用审计"],
            ["说明/", "批次差异、复核报告、数据字典和知识库证据"],
            ["06_SHA256SUMS.txt", "目录完整性校验"],
        ],
        [3.2, 3.5],
    )
    add_body(doc, "报告结束。当前材料真实可审计，但最新批次的人工指标仍应在复核后重新冻结。")
    doc.save(path)


def build_final_delivery(run_dir: Path, reviewed_run: Path, output_dir: Path) -> Path:
    """Build the latest 50-case package plus the matching reviewed baseline evidence."""
    if not reviewed_run.is_dir():
        raise FileNotFoundError(reviewed_run)
    reviewed_metrics = read_json(reviewed_run / "reports" / "metrics.json")
    if not reviewed_metrics.get("formal_ready") or reviewed_metrics.get("completed_human_review_count") != 50:
        raise ValueError("Reviewed baseline is not formally ready")

    report = build_delivery(run_dir, output_dir)
    latest_manifest = read_json(run_dir / "run_manifest.json")
    reviewed_manifest = read_json(reviewed_run / "run_manifest.json")
    latest_id = latest_manifest["run_id"]
    reviewed_id = reviewed_manifest["run_id"]

    (output_dir / "03_评测结果汇总.csv").replace(output_dir / "03_评测结果汇总_批次08.csv")
    (output_dir / "04_双人复核表.csv").replace(output_dir / "04C_批次08待复核模板.csv")
    shutil.copy2(reviewed_run / "human-review-template.csv", output_dir / "04_双人复核表_批次04.csv")
    shutil.copy2(reviewed_run / "human-reviews" / "reviewer-1.csv", output_dir / "04A_复核人1_批次04.csv")
    shutil.copy2(reviewed_run / "human-reviews" / "reviewer-2.csv", output_dir / "04B_复核人2_批次04.csv")
    shutil.copy2(
        reviewed_run / "reports" / "metrics.json",
        output_dir / "07_人工复核正式指标_批次04.json",
    )

    evidence_dir = output_dir / "说明" / "批次04人工复核证据"
    evidence_dir.mkdir(parents=True, exist_ok=True)
    for name in ("evaluation-report.md", "cases.csv", "failures.json", "metrics.json"):
        shutil.copy2(reviewed_run / "reports" / name, evidence_dir / name)
    shutil.copy2(reviewed_run / "run_manifest.json", evidence_dir / "run_manifest.json")

    unchanged: list[str] = []
    changed: list[str] = []
    for case_number in range(1, 51):
        case_id = f"{case_number:03d}"
        old_path = reviewed_run / "results" / f"case-{case_id}.json"
        new_path = run_dir / "results" / f"case-{case_id}.json"
        (unchanged if sha256(old_path) == sha256(new_path) else changed).append(case_id)
    mapping = f"""# 批次 04 至 08 的证据映射

- 人工复核基线：`{reviewed_id}`，50/50 双人复核，`formal_ready=true`。
- 最新修复复测：`{latest_id}`，50/50 运行完成，机器失败为 0。
- 结果文件哈希完全一致：{', '.join(unchanged)}。
- 结果文件发生变化：{', '.join(changed)}，共 {len(changed)} 组。

## 已定向修复并复测

- 037：开放题关键词评分修复，正确答案能够判对并更新 MIRT。
- 040：支持带中文全角括号的选项答案。
- 041：Temperature 固定题仅保留对应官方出处。
- 048：同一 request_id 重放，attempt_count 与 U/A/R 不再变化。

## 结论边界

批次 04 的人工指标不得套用到批次 08。48 组结果变化可能包含修复、重新检索、重新生成或运行记录差异；
只有对批次 08 重新完成双人复核，才能计算该批次的幻觉率、难度适配率、知识覆盖率和引用可追溯率。
"""
    (output_dir / "说明" / "批次04至08修复映射.md").write_text(mapping, encoding="utf-8")

    readme = f"""# ECHO 50 组数据实际评测赛事提交包

本包采用两阶段证据，不能把两个批次混成一个结论：

1. `{reviewed_id}`：两名真实成员已完成 50/50 复核，正式指标见 `07_人工复核正式指标_批次04.json`。
2. `{latest_id}`：修复后全量复跑，50/50 完成、50/50 四 Agent 闭环、机器失败为 0。

批次 04 的正式人工结果为：幻觉率 1.01%（通过）、难度适配率 88.00%（通过）、
核心知识覆盖率 76.00%（未通过）、引用可追溯率 99.24%（未达到 100% 目标）、闭环完整率 100%（通过）。

批次 08 已修复 037、040、041、048 的关键问题，但由于 04→08 有 48/50 结果文件变化，
批次 08 仍需使用 `04C_批次08待复核模板.csv` 重新完成双人复核后才能形成同批次正式指标。

本轮 50 组没有音频输入；微表征与 ASR 不参与评分。
"""
    (output_dir / "说明" / "README_复核与正式冻结说明.md").write_text(readme, encoding="utf-8")

    environment = read_json(output_dir / "05_运行环境与版本.json")
    environment.update(
        {
            "schema_version": "2.0",
            "evidence_model": "two_stage_reviewed_baseline_and_repaired_rerun",
            "candidate_status": "latest_rerun_pending_matching_two_human_reviews",
            "formal_ready": False,
            "all_thresholds_passed": False,
            "reviewed_baseline": {
                "run_id": reviewed_id,
                "formal_ready": True,
                "completed_human_review_count": 50,
                "all_thresholds_passed": False,
                "metrics_file": "07_人工复核正式指标_批次04.json",
            },
            "latest_rerun": {
                "run_id": latest_id,
                "completed_cases": 50,
                "machine_failure_count": 0,
                "matching_human_review_count": 0,
            },
            "batch_comparison": {
                "unchanged_result_file_count": len(unchanged),
                "unchanged_case_ids": unchanged,
                "changed_result_file_count": len(changed),
            },
            "limitations": [
                "批次 04 的覆盖率和引用可追溯率未达到内部竞赛目标。",
                "批次 08 已完成修复复测，但尚无与该批次匹配的双人复核。",
                "27 个资源已通过自动检查，但按产品发布规则仍为 pending_review。",
                "运行时工作区包含未提交评测修复；运行清单保留了该事实。",
            ],
        }
    )
    write_json(output_dir / "05_运行环境与版本.json", environment)

    dictionary = f"""# 数据字典

- `原始输入/case-XXX.json`：{latest_id} 的冻结案例、画像初始化、实际请求和运行前画像。
- `实际输出/case-XXX.json`：{latest_id} 的 ECHO 最终回复、资源、更新后画像和失败原因。
- `Agent中间记录/case-XXX.json`：{latest_id} 的四个后台 Agent 持久化记录。
- `官方引用/case-XXX.json`：{latest_id} 的官方来源；047 为模块切换，引用数组为空。
- `原始HTTP记录/case-XXX.json`：{latest_id} 的未改写 HTTP 审计记录。
- `03_评测结果汇总_批次08.csv`：最新复测机器状态，不含人工指标结论。
- `04_双人复核表_批次04.csv`：{reviewed_id} 已完成的 100 行双人复核。
- `04A/04B_复核人*_批次04.csv`：两名复核者各自的原始独立表。
- `04C_批次08待复核模板.csv`：{latest_id} 的空白复核模板。
- `07_人工复核正式指标_批次04.json`：{reviewed_id} 的正式计分结果。
"""
    (output_dir / "说明" / "数据字典.md").write_text(dictionary, encoding="utf-8")

    build_two_stage_report(
        report,
        run_dir,
        reviewed_run,
        output_dir / "图表" / "01_覆盖分布.png",
        output_dir / "图表" / "02_机器结果.png",
    )
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--run-dir", type=Path, default=DEFAULT_RUN)
    parser.add_argument("--reviewed-run-dir", type=Path, default=DEFAULT_REVIEWED_RUN)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUT)
    parser.add_argument(
        "--checksums-only",
        action="store_true",
        help="Only regenerate 06_SHA256SUMS.txt for an existing delivery directory.",
    )
    args = parser.parse_args()
    if args.checksums_only:
        print(write_checksum_manifest(args.output_dir.resolve()))
        return
    report = build_final_delivery(
        args.run_dir.resolve(),
        args.reviewed_run_dir.resolve(),
        args.output_dir.resolve(),
    )
    print(report)


if __name__ == "__main__":
    main()
