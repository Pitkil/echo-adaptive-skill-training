"""Extract fixed quiz items from structured PDF, DOCX, TXT, or Markdown files."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from config import AIConfig
from docx import Document
from openai import OpenAI
from pypdf import PdfReader

SUPPORTED_IMPORT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

FIELD_ALIASES = {
    "题目": "content",
    "题干": "content",
    "content": "content",
    "答案": "answer",
    "answer": "answer",
    "题型": "type",
    "type": "type",
    "用途": "purpose",
    "purpose": "purpose",
    "难度": "difficulty",
    "difficulty": "difficulty",
    "评分方法": "scoring_method",
    "评分标准": "scoring_method",
    "scoring_method": "scoring_method",
    "资料名称": "source_title",
    "来源名称": "source_title",
    "source_title": "source_title",
    "官方链接": "source_url",
    "来源链接": "source_url",
    "source_url": "source_url",
    "出处章节": "source_section",
    "来源章节": "source_section",
    "source_section": "source_section",
    "更新MIRT": "counts_for_mirt",
    "是否更新MIRT": "counts_for_mirt",
    "counts_for_mirt": "counts_for_mirt",
}

PURPOSE_ALIASES = {
    "前测": "pretest",
    "pretest": "pretest",
    "后测": "posttest",
    "posttest": "posttest",
    "阶段测试": "stage_test",
    "阶段测验": "stage_test",
    "stage_test": "stage_test",
    "练习": "practice",
    "practice": "practice",
}

DIFFICULTY_ALIASES = {
    "基础": "foundation",
    "foundation": "foundation",
    "标准": "standard",
    "中等": "standard",
    "standard": "standard",
    "进阶": "advanced",
    "困难": "advanced",
    "advanced": "advanced",
}


def extract_text_from_document(path: Path) -> str:
    extension = path.suffix.lower()
    if extension not in SUPPORTED_IMPORT_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_IMPORT_EXTENSIONS))
        raise ValueError(f"不支持的题库格式，当前支持：{supported}")
    if extension == ".pdf":
        return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)
    if extension == ".docx":
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            parts.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(parts)
    return path.read_text(encoding="utf-8-sig")


def _json_payload(text_value: str) -> list[dict[str, Any]]:
    candidates = [text_value.strip()]
    candidates.extend(
        match.group(1).strip()
        for match in re.finditer(r"```(?:json)?\s*([\s\S]*?)```", text_value, re.IGNORECASE)
    )
    for candidate in candidates:
        try:
            payload = json.loads(candidate)
        except json.JSONDecodeError:
            continue
        if isinstance(payload, dict):
            payload = next(
                (
                    payload[key]
                    for key in ("questions", "quizzes", "items")
                    if isinstance(payload.get(key), list)
                ),
                [payload],
            )
        if isinstance(payload, list):
            return [item for item in payload if isinstance(item, dict)]
    return []


def _labelled_payload(text_value: str) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    current: dict[str, Any] = {}
    active_field: str | None = None

    for raw_line in text_value.splitlines():
        line = raw_line.strip()
        if not line or re.fullmatch(r"-{3,}", line):
            if current.get("content") and current.get("answer"):
                items.append(current)
                current = {}
                active_field = None
            continue
        labelled = re.match(r"^(?:[-*]\s*)?([^:：]{1,20})[:：]\s*(.*)$", line)
        if labelled:
            alias = FIELD_ALIASES.get(labelled.group(1).strip())
            if alias == "content" and current.get("content") and current.get("answer"):
                items.append(current)
                current = {}
            if alias:
                current[alias] = labelled.group(2).strip()
                active_field = alias
                continue
        if active_field:
            current[active_field] = f"{current.get(active_field, '')}\n{line}".strip()

    if current.get("content") and current.get("answer"):
        items.append(current)
    return items


def _ai_payload(text_value: str) -> list[dict[str, Any]]:
    if not AIConfig.API_KEY or not AIConfig.MODEL_NAME or AIConfig.API_KEY == "test-key":
        return []
    prompt = f"""从下面题库文档中提取题目，只返回 JSON 数组。
每项字段必须包含 content、answer、type、purpose、difficulty、scoring_method、
source_title、source_url、source_section、counts_for_mirt。
purpose 只能是 pretest、posttest、stage_test、practice；
difficulty 只能是 foundation、standard、advanced。
没有依据的字段保留空字符串，不得编造官方来源。

文档内容：
{text_value[:30000]}
"""
    response = OpenAI(api_key=AIConfig.API_KEY, base_url=AIConfig.BASE_URL).chat.completions.create(
        model=AIConfig.MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_tokens=AIConfig.AGENT_MAX_TOKENS,
        response_format={"type": "json_object"},
    )
    return _json_payload(response.choices[0].message.content or "")


def _as_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value or "").strip().lower() not in {"false", "0", "否", "不更新", "no"}


def normalize_quiz_item(item: dict[str, Any], default_source_title: str) -> dict[str, Any]:
    normalized = {
        FIELD_ALIASES.get(str(key).strip(), str(key).strip()): value
        for key, value in item.items()
    }
    purpose_raw = str(normalized.get("purpose") or "practice").strip().lower()
    difficulty_raw = str(normalized.get("difficulty") or "standard").strip().lower()
    return {
        "content": str(normalized.get("content") or "").strip(),
        "answer": str(normalized.get("answer") or "").strip(),
        "type": str(normalized.get("type") or "Open").strip(),
        "purpose": PURPOSE_ALIASES.get(purpose_raw, "practice"),
        "difficulty": DIFFICULTY_ALIASES.get(difficulty_raw, "standard"),
        "scoring_method": str(normalized.get("scoring_method") or "").strip(),
        "source_title": str(normalized.get("source_title") or default_source_title).strip(),
        "source_url": str(normalized.get("source_url") or "").strip(),
        "source_section": str(normalized.get("source_section") or "").strip(),
        "counts_for_mirt": _as_bool(normalized.get("counts_for_mirt", True)),
    }


def validate_quiz_item(item: dict[str, Any]) -> list[str]:
    labels = {
        "content": "题目",
        "answer": "答案",
        "scoring_method": "评分方法",
        "source_title": "资料名称",
        "source_url": "官方链接",
        "source_section": "出处章节",
    }
    return [f"缺少{label}" for field, label in labels.items() if not str(item.get(field) or "").strip()]


def extract_quiz_preview(path: Path) -> tuple[int, list[dict[str, Any]]]:
    text_value = extract_text_from_document(path)
    if not text_value.strip():
        raise ValueError("未能从题库文件中读取文字")
    raw_items = _json_payload(text_value) or _labelled_payload(text_value) or _ai_payload(text_value)
    if not raw_items:
        raise ValueError("未识别到题目，请使用题库模板或配置出题模型后重试")
    items = [normalize_quiz_item(item, path.name) for item in raw_items]
    return len(text_value), [
        {**item, "issues": validate_quiz_item(item), "valid": not validate_quiz_item(item)}
        for item in items
    ]
