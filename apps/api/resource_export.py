"""Create polished learner-facing Word documents for generated resources."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

INK = RGBColor(29, 50, 70)
MUTED = RGBColor(89, 110, 122)
ACCENT = RGBColor(66, 111, 98)
LIGHT_BLUE = "EAF3F7"
LINE = "D8E4E9"
BODY_FONT = "Arial"
CJK_FONT = "Microsoft YaHei"


def _set_run_font(
    run,
    *,
    size: float = 10.5,
    color: RGBColor = INK,
    bold: bool | None = None,
) -> None:
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def _shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    if border:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        borders.append(left)
        properties.append(borders)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    number = OxmlElement("w:t")
    number.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    field_run._r.extend([begin, instruction, separate, number, end])
    _set_run_font(field_run, size=8.5, color=MUTED)
    tail = paragraph.add_run(" 页")
    _set_run_font(tail, size=8.5, color=MUTED)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    # A4 is the expected print format for the Chinese training audience.
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.45

    for style_name, size, before, after in (
        ("Heading 1", 15, 18, 8),
        ("Heading 2", 12, 13, 6),
    ):
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_style = document.styles["List Number"]
    list_style.font.name = BODY_FONT
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    list_style.font.size = Pt(10.5)
    list_style.paragraph_format.left_indent = Cm(0.74)
    list_style.paragraph_format.first_line_indent = Cm(-0.45)
    list_style.paragraph_format.space_after = Pt(6)
    list_style.paragraph_format.line_spacing = 1.4

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(4)
    run = header.add_run("ECHO  ·  个性化学习资源")
    _set_run_font(run, size=8.5, color=MUTED, bold=True)
    header_properties = header._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), LINE)
    borders.append(bottom)
    header_properties.append(borders)
    _add_page_number(section.footer.paragraphs[0])


def _add_labeled_line(document: Document, text: str) -> None:
    match = re.match(r"^([^：:]{1,18}[：:])\s*(.*)$", text)
    paragraph = document.add_paragraph()
    if match:
        label = paragraph.add_run(match.group(1))
        _set_run_font(label, bold=True, color=INK)
        value = paragraph.add_run(match.group(2))
        _set_run_font(value, color=INK)
    else:
        _set_run_font(paragraph.add_run(text), color=INK)


def _add_content(document: Document, content: str) -> None:
    document.add_heading("学习内容", level=1)
    for raw_line in content.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        numbered = re.match(r"^\d+[.、]\s*(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            _set_run_font(paragraph.add_run(numbered.group(1)), color=INK)
        else:
            _add_labeled_line(document, line)


def _add_callout(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.right_indent = Cm(0.25)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.35
    _shade_paragraph(paragraph, LIGHT_BLUE, "426F62")
    label_run = paragraph.add_run(f"{label}\n")
    _set_run_font(label_run, size=9, color=ACCENT, bold=True)
    value_run = paragraph.add_run(value.strip() or "根据当前学习记录生成。")
    _set_run_font(value_run, size=10, color=INK)


def _source_text(source: dict[str, Any], index: int) -> str:
    title = source.get("source_title") or source.get("title") or source.get("document_title")
    section = source.get("source_section") or source.get("section") or source.get("chapter")
    url = source.get("source_url") or source.get("url") or source.get("link")
    parts = [str(title or f"官方资料 {index}")]
    if section:
        parts.append(str(section))
    if url:
        parts.append(str(url))
    return " · ".join(parts)


def build_resource_docx(resource: Any) -> bytes:
    """Return a valid, polished DOCX for one learner-owned resource row."""

    document = Document()
    _configure_document(document)
    document.core_properties.title = resource.title
    document.core_properties.subject = "ECHO 个性化学习资源"
    document.core_properties.author = "ECHO"

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(8)
    kicker_run = kicker.add_run(
        {
            "custom_note": "定制学习资料",
            "practice_guide": "实操指南",
            "staged_test": "阶段练习",
        }.get(resource.resource_type, "学习资源")
    )
    _set_run_font(kicker_run, size=9, color=ACCENT, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    _set_run_font(title.add_run(resource.title), size=25, color=INK, bold=True)

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(18)
    difficulty = {"foundation": "基础", "standard": "标准", "advanced": "进阶"}.get(
        resource.difficulty, resource.difficulty
    )
    _set_run_font(meta.add_run(f"推荐难度  {difficulty}"), size=9, color=MUTED)

    _add_callout(document, "为什么为你推荐", resource.personalization_reason)
    _add_content(document, resource.content)

    document.add_heading("官方出处", level=1)
    sources = resource.evidence_sources or []
    if sources:
        for index, source in enumerate(sources, start=1):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.82)
            paragraph.paragraph_format.first_line_indent = Cm(-0.82)
            paragraph.paragraph_format.space_after = Pt(6)
            _set_run_font(
                paragraph.add_run(f"来源 {index}  "),
                size=9.5,
                color=ACCENT,
                bold=True,
            )
            _set_run_font(
                paragraph.add_run(_source_text(source, index)),
                size=9.5,
                color=MUTED,
            )
    else:
        paragraph = document.add_paragraph()
        _set_run_font(
            paragraph.add_run("当前资源尚未取得可追溯的官方依据，后续补充后会更新校验结果。"),
            size=9.5,
            color=MUTED,
        )

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(14)
    note.paragraph_format.space_after = Pt(0)
    _shade_paragraph(note, "F4F7F8")
    _set_run_font(
        note.add_run("本资料由 ECHO 根据当前学习记录生成。请结合课程内容与官方出处学习。"),
        size=8.5,
        color=MUTED,
    )

    # Avoid an accidental blank section created by third-party templates.
    if len(document.sections) > 1:
        document.sections[-1].start_type = WD_SECTION.CONTINUOUS
    output = BytesIO()
    document.save(output)
    return output.getvalue()
